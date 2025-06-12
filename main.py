from langchain_community.document_loaders import PDFPlumberLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_core.documents import Document
import nltk
from nltk.tokenize import sent_tokenize
import zipfile
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from embedding import embed_and_store_with_faiss
from dotenv import load_dotenv
from embedding import search_faiss, load_faiss_index, hybrid_rerank_with_cross_encoder
from rag import generate_answer_from_chunks


load_dotenv()
api_key = os.getenv("OPEN_AI_KEY")

nltk_data_path = os.path.join(os.getcwd(), "nltk_data")
if not os.path.exists(nltk_data_path):
    os.makedirs(nltk_data_path)
    nltk.download('punkt', download_dir=nltk_data_path)
    nltk.download('punkt_tab', download_dir=nltk_data_path)
    zip_path = os.path.join(nltk_data_path, 'tokenizers', 'punkt' , 'punkt.zip')
    extract_to = os.path.join(nltk_data_path, 'tokenizers', 'punkt')
    if os.path.exists(zip_path):
        os.makedirs(extract_to, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_to)

if nltk_data_path not in nltk.data.path:
    nltk.data.path.insert(0, nltk_data_path)



session='semantic' #'semantic' or 'text'

def extractor(file_path):
    if os.path.exists(file_path):
        if file_path.lower().endswith(".pdf"):

            loader = PDFPlumberLoader(file_path) # we load the file
            pages = loader.load() # we load it into documnent object

            print("Sucessfully loaded",len(pages), "from",file_path)
            return pages
        elif file_path.lower().endswith(".docx"):
            text = extract_text_and_tables_in_order(file_path)
            return text
    else:
        print("Path does not exist")


'''Adding new functions so that we can read table content from .docx files'''
def iter_block_items(parent):
    """
    Generates a reference to each paragraph and table child in document order.
    `parent` would most commonly be a reference to a main Document object,
    but also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    # This part handles different types of parents (Document, Cell, etc.)
    if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
        parent_elm = parent.element.body
    elif hasattr(parent, '_tc'): # For a _Cell object
        parent_elm = parent._tc
    elif hasattr(parent, '_tr'): # For a _Row object (though usually you iterate cells)
        parent_elm = parent._tr
    else:
        # Fallback for unexpected types, though often you'd only call this on Document or Cell
        # For simplicity, if you only expect Document objects, you can simplify this.
        if hasattr(parent, 'element'):
            parent_elm = parent.element
        else:
            raise ValueError(f"Unsupported parent type for iteration: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P): # Check if the XML element is a paragraph
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl): # Check if the XML element is a table
            yield DocxTable(child, parent)

def extract_text_and_tables_in_order(docx_path):
    doc = DocxDocument(docx_path)
    full_text_content = ""

    for block in iter_block_items(doc):
        if isinstance(block, DocxParagraph):
            if block.text.strip():
                full_text_content += block.text + "\n"
        elif isinstance(block, DocxTable):
            # Process table content
            for row in block.rows:
                row_cells_text = [cell.text.strip() for cell in row.cells]
                full_text_content += " ".join(row_cells_text) + "\n"
            full_text_content += "\n" # Add a newline to separate tables

    text_as_doc = Document(page_content=full_text_content, metadata={"source": os.path.basename(docx_path),"file_path":docx_path ,"file_type": "docx"})
    return [text_as_doc]

'''End of those newly added '''
def extract_text_from_docx(docx_path):
    doc = DocxDocument(docx_path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    text_as_doc = Document(page_content=text, metadata={"source": os.path.basename(docx_path), "file_type": "docx"})
    return [text_as_doc]


def text_chunker(pages, chunk_size = 500 , chunk_overlap = 200):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=200)
    chunks = text_splitter.split_documents(pages)
    print("Split into", len(chunks), "text chunks")
    with open("output_text.txt", "w", encoding="utf-8") as f:
        for i, chunk in enumerate(chunks):
            header = f"____Chunk{i + 1}____\n"
            content = f"{chunk.page_content[:500]}\n"
            print(header)
            print(content)
            f.write(header)
            f.write(content)
    f.close()
    return chunks


def semantic_chunker(pages, chunk_size=500, chunk_overlap=200):
    chunks = []
    for doc_index, doc in enumerate(pages):
        sentences = sent_tokenize(doc.page_content)
        current_sentences_for_chunk = []

        for sentence_index, sentence_text_from_doc in enumerate(sentences):
            candidate_sentences = current_sentences_for_chunk + [sentence_text_from_doc]
            candidate_joined_text = " ".join(candidate_sentences)

            if current_sentences_for_chunk and len(candidate_joined_text) > chunk_size:
                emitted_chunk_text = candidate_joined_text
                emitted_chunk_sentence_list = candidate_sentences
                chunk_metadata = {
                    "source": doc.metadata.get("source"),
                    "file_path": doc.metadata.get("file_path"),
                    "page_number": doc_index + 1,
                    "chunk_index": len(chunks) + 1
                }
                chunks.append(Document(page_content=emitted_chunk_text, metadata=chunk_metadata))

                new_base_for_next_chunk_parts = []
                if not emitted_chunk_sentence_list:
                    current_sentences_for_chunk = []
                else:
                    last_sentence_in_emitted_chunk = emitted_chunk_sentence_list[-1]
                    if len(last_sentence_in_emitted_chunk) > chunk_overlap:
                        overlap_content = emitted_chunk_text[-chunk_overlap:]
                        new_base_for_next_chunk_parts = [overlap_content]
                    else:
                        sentence_based_overlap_list = []
                        for s_overlap in reversed(emitted_chunk_sentence_list):
                            temp_sentences = [s_overlap] + sentence_based_overlap_list
                            temp_joined_text = " ".join(temp_sentences)

                            if len(temp_joined_text) >= chunk_overlap:
                                if sentence_based_overlap_list and len(
                                        " ".join(sentence_based_overlap_list)) >= chunk_overlap:
                                    pass
                                else:
                                    sentence_based_overlap_list.insert(0, s_overlap)
                                break
                            else:
                                sentence_based_overlap_list.insert(0, s_overlap)

                            if len(sentence_based_overlap_list) == len(emitted_chunk_sentence_list) and len(
                                    temp_joined_text) < chunk_overlap:
                                break
                        new_base_for_next_chunk_parts = sentence_based_overlap_list
                current_sentences_for_chunk = new_base_for_next_chunk_parts
            else:
                current_sentences_for_chunk.append(sentence_text_from_doc)

        if current_sentences_for_chunk:
            final_chunk_text = " ".join(current_sentences_for_chunk)
            chunk_metadata = {
                "source": doc.metadata.get("source"),
                "file_path": doc.metadata.get("file_path"),
                "page_number": doc_index + 1,
                "chunk_index": len(chunks) + 1
            }
            chunks.append(Document(page_content=final_chunk_text, metadata=chunk_metadata))

    output_filename = "output_semantic.txt"
    try:
        with open(output_filename, "w", encoding="utf-8") as f:
            for i, chunk_doc in enumerate(chunks):
                header = f"____Chunk{i + 1}____\n"
                page_content_to_write = chunk_doc.page_content
                content_display = page_content_to_write[:chunk_size]
                content_display_full = content_display
                if len(page_content_to_write) > chunk_size:
                    content_display_full += f"\n... (actual length: {len(page_content_to_write)}, chunk_size: {chunk_size})"
                f.write(header)
                f.write(content_display_full + "\n")
        # print(f"Production chunking output written to {output_filename}") # Optional: uncomment if needed
    except Exception as e:
        # In production, consider proper logging instead of print
        # print(f"Error writing production output file: {e}")
        pass  # Or raise e, or log e

    return chunks


if __name__ == "__main__":
    task = input("Enter what to do - embed or query ")
    if task == "embed":
        file_path = input("Enter the name of the path")
        if session=='text':
            pages = extractor(file_path)
            print(pages)
            chunks = text_chunker(pages, 500, 200)

        elif session=='semantic':
            pages = extractor(file_path)
            output_parsed_text = "output_parsed_text"
            with open(output_parsed_text, "w", encoding="utf-8") as f:
                for i, doc_item in enumerate(pages):
                        # Correct way to access and write page content
                        f.write(f"--- Page {i+1} ---\n")
                        f.write(doc_item.page_content)
                        f.write("\n\n") # Add some separation between pages
            print(f"Number of documents (pages) extracted: {len(pages)}")
            for i, doc_item in enumerate(pages):
                print(f"Content of doc {i} (first 150 chars): {doc_item.page_content[:150]}")
            chunks = semantic_chunker(pages, 5000, 800)
            faiss_index = embed_and_store_with_faiss(
                chunks=chunks,
                openai_api_key=api_key,
                save_path="faiss_index_store"
            )
    elif task == "query":
        faiss_index = load_faiss_index('faiss_index_store', openai_api_key=api_key)
        query = input("Enter query - ")
        candidates = faiss_index.similarity_search(query, k=20)
        results = hybrid_rerank_with_cross_encoder(query, candidates, top_k=3)
        answer = generate_answer_from_chunks(query, results, api_key)
        print(answer)


