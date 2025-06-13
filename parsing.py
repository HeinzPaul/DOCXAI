from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
import os


def extractor(file_path):
    if os.path.exists(file_path):
        if file_path.lower().endswith(".pdf"):

            loader = PDFPlumberLoader(file_path) # we load the file
            pages = loader.load() # we load it into documnent object

            print("Sucessfully loaded",len(pages), "from",file_path)
            return pages
        elif file_path.lower().endswith(".docx"):
            text = extract_text_and_tables_in_order(file_path)
            return text
    else:
        print("Path does not exist")


def iter_block_items(parent):
    if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
        parent_elm = parent.element.body
    elif hasattr(parent, '_tc'): # For a _Cell object
        parent_elm = parent._tc
    elif hasattr(parent, '_tr'): # For a _Row object (though usually you iterate cells)
        parent_elm = parent._tr
    else:
        # Fallback for unexpected types, though often you'd only call this on Document or Cell
        # For simplicity, if you only expect Document objects, you can simplify this.
        if hasattr(parent, 'element'):
            parent_elm = parent.element
        else:
            raise ValueError(f"Unsupported parent type for iteration: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P): # Check if the XML element is a paragraph
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl): # Check if the XML element is a table
            yield DocxTable(child, parent)

def extract_text_and_tables_in_order(docx_path):
    doc = DocxDocument(docx_path)
    full_text_content = ""

    for block in iter_block_items(doc):
        if isinstance(block, DocxParagraph):
            if block.text.strip():
                full_text_content += block.text + "\n"
        elif isinstance(block, DocxTable):
            # Process table content
            for row in block.rows:
                row_cells_text = [cell.text.strip() for cell in row.cells]
                full_text_content += " ".join(row_cells_text) + "\n"
            full_text_content += "\n" # Add a newline to separate tables

    text_as_doc = Document(page_content=full_text_content, metadata={"source": os.path.basename(docx_path),"file_path":docx_path ,"file_type": "docx"})
    return [text_as_doc]

def extract_text_from_docx(docx_path):
    doc = DocxDocument(docx_path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    text_as_doc = Document(page_content=text, metadata={"source": os.path.basename(docx_path), "file_type": "docx"})
    return [text_as_doc]